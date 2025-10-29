from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
import logging
import re
from datetime import datetime

# 导入RAG导出模块
from .rag_export import RAGExporter, export_for_rag_knowledge_base

# 设置日志
logger = logging.getLogger(__name__)

def sanitize_filename(filename):
    """清理文件名，移除不合法字符"""
    # 移除Windows不支持的字符
    illegal_chars = r'<>:"/\|?*'
    for char in illegal_chars:
        filename = filename.replace(char, '_')
    # 移除连续空格
    filename = re.sub(r'\s+', '_', filename)
    # 限制文件名长度
    if len(filename) > 200:
        filename = filename[:200]
    return filename

def export_to_word(data, file_path):
    """导出政策数据到Word文档"""
    doc = Document()
    
    # 设置中文字体 - 使用更安全的方式
    try:
        # 直接设置段落字体，避免样式设置问题
        pass
    except Exception:
        # 如果字体设置失败，继续执行，不影响文档生成
        pass
    
    # 添加标题
    title = doc.add_heading('空间规划政策汇总', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加统计信息
    stats_para = doc.add_paragraph(f'共导出 {len(data)} 条政策文件')
    stats_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加目录
    toc_heading = doc.add_heading('目录', level=1)
    toc_para = doc.add_paragraph()
    
    # 生成目录内容
    for i, policy in enumerate(data, 1):
        # 解析政策数据格式
        if isinstance(policy, (list, tuple)):
            title = str(policy[2]) if len(policy) > 2 else "未知标题"
            pub_date = str(policy[3]) if len(policy) > 3 else "未知日期"
        elif isinstance(policy, dict):
            title = str(policy.get('title', '未知标题'))
            pub_date = str(policy.get('pub_date', '未知日期'))
        else:
            title = "未知"
            pub_date = "未知日期"
        
        # 添加目录项（包含时间）
        toc_para.add_run(f'{i}. {title} ({pub_date})\n')
    
    # 添加分隔线
    doc.add_paragraph('=' * 50)
    
    # 为每个政策添加内容
    for i, policy in enumerate(data, 1):
        # 解析政策数据格式
        if isinstance(policy, (list, tuple)):
            title = str(policy[2]) if len(policy) > 2 else "未知标题"
            level = str(policy[1]) if len(policy) > 1 else "未知层级"
            pub_date = str(policy[3]) if len(policy) > 3 else "未知日期"
            source = str(policy[4]) if len(policy) > 4 else "未知来源"
            content = str(policy[5]) if len(policy) > 5 else "无内容"
        elif isinstance(policy, dict):
            title = str(policy.get('title', '未知标题'))
            level = str(policy.get('level', '未知层级'))
            pub_date = str(policy.get('pub_date', '未知日期'))
            source = str(policy.get('source', '未知来源'))
            content = str(policy.get('content', '无内容'))
        else:
            title = level = pub_date = source = content = "未知"
        
        # 政策标题
        policy_title = doc.add_heading(f'{i}. {title}', level=1)
        
        # 基本信息
        info_para = doc.add_paragraph()
        info_para.add_run(f'层级：{level}\n')
        info_para.add_run(f'发布日期：{pub_date}\n')
        info_para.add_run(f'来源：{source}\n')
        
        # 正文内容
        content_para = doc.add_paragraph()
        content_para.add_run('正文：\n')
        content_para.add_run(content)
        
        # 添加分隔线
        if i < len(data):
            doc.add_paragraph('=' * 50)
    
    # 保存文档
    doc.save(file_path)
    return True

class DataExporter:
    """数据导出器"""
    
    def __init__(self):
        pass
    
    def export_to_word(self, data, file_path):
        """导出政策数据到Word文档"""
        return export_to_word(data, file_path)
    
    def export_to_excel(self, data, file_path):
        """导出政策数据到Excel文档"""
        # 检查pandas是否可用
        try:
            import pandas as pd  # type: ignore
        except ImportError:
            print("❌ 缺少依赖库: pandas")
            print("💡 解决方案: pip install pandas openpyxl")
            return False
        
        try:
            # 创建Excel工作簿
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment
            
            wb = Workbook()
            
            # 创建目录工作表
            toc_sheet = wb.active
            if toc_sheet is not None:
                toc_sheet.title = "目录"
                
                # 设置目录标题
                toc_sheet['A1'] = "空间规划政策汇总目录"
                toc_sheet['A1'].font = Font(bold=True, size=16)
                toc_sheet['A1'].alignment = Alignment(horizontal='center')
                toc_sheet.merge_cells('A1:D1')
                
                # 添加统计信息
                toc_sheet['A2'] = f"共导出 {len(data)} 条政策文件"
                toc_sheet['A2'].alignment = Alignment(horizontal='center')
                toc_sheet.merge_cells('A2:D2')
                
                # 添加目录内容
                toc_sheet['A4'] = "序号"
                toc_sheet['B4'] = "政策标题"
                toc_sheet['C4'] = "发布日期"
                toc_sheet['D4'] = "层级"
                toc_sheet['A4'].font = Font(bold=True)
                toc_sheet['B4'].font = Font(bold=True)
                toc_sheet['C4'].font = Font(bold=True)
                toc_sheet['D4'].font = Font(bold=True)
                
                for i, policy in enumerate(data, 1):
                    # 解析政策数据格式
                    if isinstance(policy, (list, tuple)):
                        title = str(policy[2]) if len(policy) > 2 else "未知标题"
                        pub_date = str(policy[3]) if len(policy) > 3 else "未知日期"
                        level = str(policy[1]) if len(policy) > 1 else "未知层级"
                    elif isinstance(policy, dict):
                        title = str(policy.get('title', '未知标题'))
                        pub_date = str(policy.get('pub_date', '未知日期'))
                        level = str(policy.get('level', '未知层级'))
                    else:
                        title = pub_date = level = "未知"
                    
                    toc_sheet[f'A{i+4}'] = i
                    toc_sheet[f'B{i+4}'] = title
                    toc_sheet[f'C{i+4}'] = pub_date
                    toc_sheet[f'D{i+4}'] = level
                
                # 调整列宽
                toc_sheet.column_dimensions['A'].width = 8
                toc_sheet.column_dimensions['B'].width = 50
                toc_sheet.column_dimensions['C'].width = 15
                toc_sheet.column_dimensions['D'].width = 15
            
            # 创建详细内容工作表
            detail_sheet = wb.create_sheet("详细内容")
            
            # 转换数据格式
            df_data = []
            for policy in data:
                # 解析政策数据格式
                if isinstance(policy, (list, tuple)):
                    policy_id = str(policy[0]) if len(policy) > 0 else ""
                    level = str(policy[1]) if len(policy) > 1 else "未知层级"
                    title = str(policy[2]) if len(policy) > 2 else "未知标题"
                    pub_date = str(policy[3]) if len(policy) > 3 else "未知日期"
                    source = str(policy[4]) if len(policy) > 4 else "未知来源"
                    content = str(policy[5]) if len(policy) > 5 else "无内容"
                elif isinstance(policy, dict):
                    policy_id = str(policy.get('id', ''))
                    level = str(policy.get('level', '未知层级'))
                    title = str(policy.get('title', '未知标题'))
                    pub_date = str(policy.get('pub_date', '未知日期'))
                    source = str(policy.get('source', '未知来源'))
                    content = str(policy.get('content', '无内容'))
                else:
                    policy_id = level = title = pub_date = source = content = "未知"
                
                df_data.append({
                    'ID': policy_id,
                    '层级': level,
                    '标题': title,
                    '发布日期': pub_date,
                    '来源': source,
                    '内容': content
                })
            
            # 将详细内容写入第二个工作表
            df = pd.DataFrame(df_data)
            for row_idx, (index, row) in enumerate(df.iterrows(), 1):
                for col_idx, (col_name, value) in enumerate(row.items(), 1):
                    detail_sheet.cell(row=row_idx+1, column=col_idx, value=value)
            
            # 设置详细内容表头
            for col_idx, col_name in enumerate(df.columns, 1):
                detail_sheet.cell(row=1, column=col_idx, value=col_name).font = Font(bold=True)
            
            # 调整详细内容列宽
            detail_sheet.column_dimensions['A'].width = 8
            detail_sheet.column_dimensions['B'].width = 15
            detail_sheet.column_dimensions['C'].width = 60
            detail_sheet.column_dimensions['D'].width = 15
            detail_sheet.column_dimensions['E'].width = 40
            detail_sheet.column_dimensions['F'].width = 80
            
            wb.save(file_path)
            return True
        except Exception as e:
            print(f"导出Excel文件失败: {e}")
            return False
    
    def export_to_txt(self, data, file_path):
        """导出政策数据到文本文件"""
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write('空间规划政策汇总\n')
                f.write('=' * 50 + '\n')
                f.write(f'共导出 {len(data)} 条政策文件\n\n')
                
                # 添加目录
                f.write('目录\n')
                f.write('-' * 30 + '\n')
                for i, policy in enumerate(data, 1):
                    # 解析政策数据格式
                    if isinstance(policy, (list, tuple)):
                        title = str(policy[2]) if len(policy) > 2 else "未知标题"
                        pub_date = str(policy[3]) if len(policy) > 3 else "未知日期"
                    elif isinstance(policy, dict):
                        title = str(policy.get('title', '未知标题'))
                        pub_date = str(policy.get('pub_date', '未知日期'))
                    else:
                        title = "未知"
                        pub_date = "未知日期"
                    
                    f.write(f'{i}. {title} ({pub_date})\n')
                
                f.write('\n' + '=' * 50 + '\n\n')
                
                # 详细内容
                for i, policy in enumerate(data, 1):
                    # 解析政策数据格式
                    if isinstance(policy, (list, tuple)):
                        title = str(policy[2]) if len(policy) > 2 else "未知标题"
                        level = str(policy[1]) if len(policy) > 1 else "未知层级"
                        pub_date = str(policy[3]) if len(policy) > 3 else "未知日期"
                        source = str(policy[4]) if len(policy) > 4 else "未知来源"
                        content = str(policy[5]) if len(policy) > 5 else "无内容"
                    elif isinstance(policy, dict):
                        title = str(policy.get('title', '未知标题'))
                        level = str(policy.get('level', '未知层级'))
                        pub_date = str(policy.get('pub_date', '未知日期'))
                        source = str(policy.get('source', '未知来源'))
                        content = str(policy.get('content', '无内容'))
                    else:
                        title = level = pub_date = source = content = "未知"
                    
                    f.write(f'{i}. {title}\n')
                    f.write(f'层级：{level}\n')
                    f.write(f'发布日期：{pub_date}\n')
                    f.write(f'来源：{source}\n')
                    f.write(f'正文：{content}\n')
                    f.write('=' * 50 + '\n\n')
            
            return True
        except Exception as e:
            print(f"导出文本文件失败: {e}")
            return False
    
    def export_to_markdown(self, data, file_path):
        """导出政策数据到Markdown文档"""
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write('# 空间规划政策汇总\n\n')
                f.write(f'**共导出 {len(data)} 条政策文件**\n\n')
                f.write('---\n\n')
                
                # 添加目录
                f.write('## 目录\n\n')
                for i, policy in enumerate(data, 1):
                    # 解析政策数据格式
                    if isinstance(policy, (list, tuple)):
                        title = str(policy[2]) if len(policy) > 2 else "未知标题"
                        pub_date = str(policy[3]) if len(policy) > 3 else "未知日期"
                    elif isinstance(policy, dict):
                        title = str(policy.get('title', '未知标题'))
                        pub_date = str(policy.get('pub_date', '未知日期'))
                    else:
                        title = "未知"
                        pub_date = "未知日期"
                    
                    f.write(f'{i}. [{title} ({pub_date})](#{i}-{title.replace(" ", "-")})\n')
                
                f.write('\n---\n\n')
                
                # 详细内容
                for i, policy in enumerate(data, 1):
                    # 解析政策数据格式
                    if isinstance(policy, (list, tuple)):
                        title = str(policy[2]) if len(policy) > 2 else "未知标题"
                        level = str(policy[1]) if len(policy) > 1 else "未知层级"
                        pub_date = str(policy[3]) if len(policy) > 3 else "未知日期"
                        source = str(policy[4]) if len(policy) > 4 else "未知来源"
                        content = str(policy[5]) if len(policy) > 5 else "无内容"
                    elif isinstance(policy, dict):
                        title = str(policy.get('title', '未知标题'))
                        level = str(policy.get('level', '未知层级'))
                        pub_date = str(policy.get('pub_date', '未知日期'))
                        source = str(policy.get('source', '未知来源'))
                        content = str(policy.get('content', '无内容'))
                    else:
                        title = level = pub_date = source = content = "未知"
                    
                    f.write(f'## {i}. {title}\n\n')
                    f.write(f'**层级：** {level}\n\n')
                    f.write(f'**发布日期：** {pub_date}\n\n')
                    f.write(f'**来源：** {source}\n\n')
                    f.write(f'**正文：**\n\n')
                    f.write(f'{content}\n\n')
                    f.write('---\n\n')
            
            return True
        except Exception as e:
            print(f"导出Markdown文件失败: {e}")
            return False
    
    def export_for_rag(self, data, output_dir, format_type='markdown', max_chunk_size=4096):
        """
        为RAG知识库导出优化格式
        
        Args:
            data: 政策数据列表
            output_dir: 输出目录
            format_type: 输出格式 ('markdown', 'json', 'txt')
            max_chunk_size: 最大段落大小
        
        Returns:
            导出结果信息
        """
        try:
            return export_for_rag_knowledge_base(data, output_dir, format_type, max_chunk_size)
        except Exception as e:
            print(f"RAG导出失败: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def export_individual_files(self, data, output_dir, format_type='word'):
        """
        分条导出 - 一个政策一个文件，所有文件直接放在指定目录
        
        Args:
            data: 政策数据列表
            output_dir: 输出目录
            format_type: 输出格式 ('word', 'txt', 'markdown')
        
        Returns:
            Dict: 导出结果信息
        """
        try:
            # 创建输出目录
            os.makedirs(output_dir, exist_ok=True)
            
            exported_files = []
            total_files = 0
            
            # 直接在该目录下导出所有文件
            for i, policy in enumerate(data, 1):
                # 解析政策数据
                if isinstance(policy, dict):
                    pub_date = policy.get('pub_date', '未知日期')
                    title = policy.get('title', '未知标题')
                    level = policy.get('level', '未知层级')
                    source = policy.get('source', '未知来源')
                    content = policy.get('content', '无内容')
                elif isinstance(policy, (list, tuple)):
                    pub_date = str(policy[3]) if len(policy) > 3 else "未知日期"
                    title = str(policy[2]) if len(policy) > 2 else "未知标题"
                    level = str(policy[1]) if len(policy) > 1 else "未知层级"
                    source = str(policy[4]) if len(policy) > 4 else "未知来源"
                    content = str(policy[5]) if len(policy) > 5 else "无内容"
                else:
                    pub_date = title = level = source = content = "未知"
                
                policy_data = {
                    'title': title,
                    'level': level,
                    'pub_date': pub_date,
                    'source': source,
                    'content': content
                }
                
                # 生成文件名
                safe_title = sanitize_filename(title)
                file_index = f"{i:04d}"  # 4位数字序号
                filename = f"{file_index}_{safe_title}"
                
                # 根据格式类型生成文件
                if format_type == 'word':
                    filepath = os.path.join(output_dir, f"{filename}.docx")
                    self._export_single_word(policy_data, filepath)
                elif format_type == 'txt':
                    filepath = os.path.join(output_dir, f"{filename}.txt")
                    self._export_single_txt(policy_data, filepath)
                elif format_type == 'markdown':
                    filepath = os.path.join(output_dir, f"{filename}.md")
                    self._export_single_markdown(policy_data, filepath)
                else:
                    continue
                
                exported_files.append(filepath)
                total_files += 1
            
            return {
                'success': True,
                'total_files': total_files,
                'files': exported_files,
                'output_dir': output_dir
            }
            
        except Exception as e:
            logger.error(f"分条导出失败: {e}")
            return {
                'success': False,
                'error': str(e),
                'total_files': 0,
                'files': []
            }
    
    def _export_single_word(self, policy, filepath):
        """导出单个政策为Word文档"""
        doc = Document()
        
        # 添加标题
        title = doc.add_heading(policy['title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加基本信息
        info_para = doc.add_paragraph()
        info_para.add_run(f'发布机构：{policy["level"]}\n')
        info_para.add_run(f'发布日期：{policy["pub_date"]}\n')
        info_para.add_run(f'来源：{policy["source"]}\n')
        
        # 添加正文
        doc.add_paragraph('正文：')
        content_para = doc.add_paragraph(policy['content'])
        
        # 保存文档
        doc.save(filepath)
    
    def _export_single_txt(self, policy, filepath):
        """导出单个政策为文本文件"""
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(f"{policy['title']}\n")
            f.write('=' * 80 + '\n\n')
            f.write(f"发布机构：{policy['level']}\n")
            f.write(f"发布日期：{policy['pub_date']}\n")
            f.write(f"来源：{policy['source']}\n\n")
            f.write('-' * 80 + '\n\n')
            f.write("正文：\n")
            f.write(policy['content'])
    
    def _export_single_markdown(self, policy, filepath):
        """导出单个政策为Markdown文档"""
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(f"# {policy['title']}\n\n")
            f.write("---\n\n")
            f.write(f"**发布机构：** {policy['level']}\n\n")
            f.write(f"**发布日期：** {policy['pub_date']}\n\n")
            f.write(f"**来源：** {policy['source']}\n\n")
            f.write("---\n\n")
            f.write("## 正文\n\n")
            f.write(policy['content'])
    
    # PDF导出功能已移除 